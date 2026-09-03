import docx
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm

doc = docx.Document('backend/assets/templates/COMMERCIAL PROPOSAL_33KV.docx')

# The document currently has 1 section.
# Its properties are at the end of doc._element.body: doc._element.body.sectPr
body = doc._element.body
doc_sectPr = body.sectPr

# 1. Create a copy of the document's sectPr for Section 1
sectPr1 = copy.deepcopy(doc_sectPr)

# 2. Insert sectPr1 into the paragraph just before the last one
# Last paragraph contains the image
p_last = doc.paragraphs[-1]._element
p_prev = doc.paragraphs[-2]._element

pPr = p_prev.get_or_add_pPr()
pPr.append(sectPr1)

# Now the document has 2 sections.
# doc.sections[-1] corresponds to doc_sectPr (Section 2)
# Let's modify doc_sectPr to remove header and set margins to 0

sec2 = doc.sections[-1]
# Set margins to 0
sec2.top_margin = Cm(0)
sec2.bottom_margin = Cm(0)
sec2.left_margin = Cm(0)
sec2.right_margin = Cm(0)

# Remove header references from doc_sectPr
for header_ref in doc_sectPr.findall(qn('w:headerReference')):
    doc_sectPr.remove(header_ref)
for footer_ref in doc_sectPr.findall(qn('w:footerReference')):
    doc_sectPr.remove(footer_ref)

doc.save('backend/assets/templates/test_output.docx')
print("Done")
