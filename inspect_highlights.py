import sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def inspect(filepath):
    doc = Document(filepath)
    highlights = set()

    def process_runs(runs):
        for r in runs:
            if r.font.highlight_color is not None and r.font.highlight_color != WD_COLOR_INDEX.AUTO:
                text = r.text.strip()
                if text:
                    highlights.add(text)

    for p in doc.paragraphs:
        process_runs(p.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_runs(p.runs)
                    
    print(f"Highlights in {filepath}:")
    for h in sorted(list(highlights)):
        print(f" - '{h}'")
    print("\n")

inspect("backend/assets/templates/technical_proposal_template.docx")
inspect("backend/assets/templates/commercial_proposal_template.docx")
