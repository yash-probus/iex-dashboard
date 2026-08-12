from docx import Document

def inject_tags(filepath):
    doc = Document(filepath)
    modified = False
    
    for i, p in enumerate(doc.paragraphs):
        # Check if paragraph has a picture
        has_pic = False
        for r in p.runs:
            if '<w:drawing>' in r._element.xml:
                has_pic = True
        
        if has_pic:
            # Check surrounding text for Picture 3 (monthly savings)
            prev_text = ""
            for j in range(max(0, i-3), i):
                prev_text += doc.paragraphs[j].text + " "
                
            if "Actual data shows a recurring reduction opportunity" in prev_text:
                # Remove runs with pictures
                for r in list(p.runs):
                    if '<w:drawing>' in r._element.xml:
                        p._element.remove(r._element)
                # Add dynamic tag
                p.add_run('<<%monthly_savings_chart>>')
                print(f"Replaced monthly savings chart in {filepath}")
                modified = True
                
            # Check for Add on features (dashboard screenshot)
            if "Add on Features-" in prev_text:
                # Remove runs with pictures
                for r in list(p.runs):
                    if '<w:drawing>' in r._element.xml:
                        p._element.remove(r._element)
                # Add dynamic tag
                p.add_run('<<%dashboard_screenshot>>')
                print(f"Replaced dashboard screenshot in {filepath}")
                modified = True

    if modified:
        doc.save(filepath)
        print(f"Saved {filepath}")

inject_tags('backend/assets/templates/technical_proposal_template.docx')
inject_tags('backend/assets/templates/commercial_proposal_template.docx')
