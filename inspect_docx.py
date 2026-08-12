import sys
from docx import Document

def inspect_and_replace_highlights(filepath, output_path):
    print(f"Inspecting: {filepath}")
    doc = Document(filepath)
    highlight_count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color:
                text = run.text.strip()
                if text:
                    print(f"Found highlighted text: '{text}' (Color: {run.font.highlight_color})")
                    # Optionally replace with a generic tag if we wanted to
                    # run.text = f"<<{text}>>"
                    # run.font.highlight_color = None
                    highlight_count += 1
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color:
                            text = run.text.strip()
                            if text:
                                print(f"Found highlighted text in table: '{text}'")
                                highlight_count += 1
                                
    print(f"Total highlights found: {highlight_count}")

if __name__ == "__main__":
    inspect_and_replace_highlights(sys.argv[1], sys.argv[2])
