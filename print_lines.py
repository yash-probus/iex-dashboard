from docx import Document

def print_lines(filepath):
    doc = Document(filepath)
    for i in range(60, 90):
        print(f"P{i}: {doc.paragraphs[i].text.strip()}")
        for r in doc.paragraphs[i].runs:
            if '<w:drawing>' in r._element.xml:
                print(f"  [PICTURE IN P{i}]")

print_lines('backend/assets/templates/technical_proposal_template.docx')
