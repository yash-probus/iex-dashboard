from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def apply_header_style(cell, text):
    cell.text = text
    shading_elm = parse_xml(r'<w:shd {} w:fill="0F204B"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

def apply_body_style(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(10)

def fix_doc(filepath):
    doc = Document(filepath)
    
    # We already replaced 'Six months' in the previous run, but it might be '<<months_count_word>> months'.
    # If starting from scratch, it will be 'Six months'.
    for p in doc.paragraphs:
        if 'Six months' in p.text:
            p.text = p.text.replace('Six months', '<<months_count_word>> months')
            for r in p.runs:
                if '<<months_count_word>>' in r.text:
                    r.font.bold = True
    
    t1 = None
    t2 = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            cell_0_text = table.rows[0].cells[0].text
            cell_1_text = table.rows[0].cells[1].text if len(table.rows[0].cells) > 1 else ""
            if 'Metric' in cell_0_text or 'Total Cleared Units' in cell_1_text:
                t1 = table
            elif 'Month' in cell_0_text and 'Cleared vs' in cell_1_text:
                t2 = table
                
    if t1:
        new_t1 = doc.add_table(rows=2, cols=5)
        new_t1.style = 'Table Grid'
        
        headers = ["Month", "Total Cleared Units@ Consumer bus", "Total Consumption As per E-bill", "Total Power Cost through Open Access", "Discom Cost"]
        for idx, h in enumerate(headers):
            apply_header_style(new_t1.rows[0].cells[idx], h)
            
        tags = ["<<#monthlyData>><<month_name>>", "<<cleared>>", "<<consumption>>", "<<oa_cost>>", "<<discom_cost>><</monthlyData>>"]
        for idx, t in enumerate(tags):
            apply_body_style(new_t1.rows[1].cells[idx], t)
            
        t1._element.addprevious(new_t1._element)
        t1._element.getparent().remove(t1._element)
        
    if t2:
        new_t2 = doc.add_table(rows=2, cols=6)
        new_t2.style = 'Table Grid'
        
        headers = ["Month", "Cleared vs Actual consumption %", "Power Purchase Cost (Discom Only)", "Power Purchase Cost (With Prolt)", "Total Saving", "Saving/Unit"]
        for idx, h in enumerate(headers):
            apply_header_style(new_t2.rows[0].cells[idx], h)
            
        tags = ["<<#monthlyData>><<month_name>>", "<<cleared_pct>>", "<<ppc_discom>>", "<<ppc_prolt>>", "<<saving>>", "<<saving_unit>><</monthlyData>>"]
        for idx, t in enumerate(tags):
            apply_body_style(new_t2.rows[1].cells[idx], t)
            
        t2._element.addprevious(new_t2._element)
        t2._element.getparent().remove(t2._element)
        
    doc.save(filepath)

import subprocess
# checkout to ensure we start from the state with "Six months"
subprocess.run(['git', 'checkout', 'b5fefe5e9aef549f497ddf6615e8752ec59f335c', '--', 'backend/assets/templates/technical_proposal_template.docx'])
subprocess.run(['git', 'checkout', 'b5fefe5e9aef549f497ddf6615e8752ec59f335c', '--', 'backend/assets/templates/commercial_proposal_template.docx'])

fix_doc('backend/assets/templates/technical_proposal_template.docx')
fix_doc('backend/assets/templates/commercial_proposal_template.docx')
