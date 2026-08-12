from docx import Document

def replace_last_picture(filepath):
    doc = Document(filepath)
    # Find all paragraphs that contain a picture
    pic_paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        has_pic = False
        for r in p.runs:
            if '<w:drawing>' in r._element.xml:
                has_pic = True
        if has_pic:
            pic_paragraphs.append(p)
            
    if pic_paragraphs:
        last_pic_p = pic_paragraphs[-1]
        last_pic_p.clear()
        last_pic_p.add_run('{%monthly_savings_chart}')
    doc.save(filepath)

def replace_95_percent(filepath):
    doc = Document(filepath)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '95%' in cell.text:
                    cell.text = cell.text.replace('95%', '<<procurement_potential>>%')
    doc.save(filepath)


replace_last_picture('backend/assets/templates/technical_proposal_template.docx')
replace_95_percent('backend/assets/templates/technical_proposal_template.docx')

replace_last_picture('backend/assets/templates/commercial_proposal_template.docx')
