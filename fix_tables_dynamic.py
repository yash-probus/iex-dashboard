from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def apply_header_style(cell, text):
    cell.text = text
    # Set background color to dark blue
    shading_elm = parse_xml(r'<w:shd {} w:fill="0F204B"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Set text color to white and bold
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

def apply_body_style(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(10)

def fix_doc(filepath):
    doc = Document(filepath)
    
    # Replace 'Six months'
    for p in doc.paragraphs:
        for r in p.runs:
            if 'Six' in r.text and 'months' in r.text:
                r.text = r.text.replace('Six months', '<<months_count_word>> months')
            elif 'Six' in r.text:
                # Might be split across runs
                pass
                
    # Also just blindly replace 'Six months' in paragraph text if runs didn't catch it
    for p in doc.paragraphs:
        if 'Six months' in p.text:
            # Rebuilding runs for this paragraph
            p.text = p.text.replace('Six months', '<<months_count_word>> months')
            # Warning: this clears formatting for the whole paragraph! 
            # But the 'Six months' is in a header paragraph. Let's make it bold.
            for r in p.runs:
                if '<<months_count_word>>' in r.text:
                    r.font.bold = True
                    # Let's add yellow highlight back? The user likes the dynamic part to work.
                    
    # Let's do it safely
    
    t1 = None
    t2 = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            if 'Metric' in table.rows[0].cells[0].text:
                t1 = table
            elif 'Month' in table.rows[0].cells[0].text and 'Cleared vs' in table.rows[0].cells[1].text:
                t2 = table
                
    if t1:
        new_t1 = doc.add_table(rows=2, cols=5)
        new_t1.style = 'Table Grid'
        
        headers = ["Month", "Total Cleared Units@ Consumer bus", "Total Consumption As per E-bill", "Total Power Cost through Open Access", "Discom Cost"]
        for idx, h in enumerate(headers):
            apply_header_style(new_t1.rows[0].cells[idx], h)
            
        tags = ["{#monthlyData}<<month_name>>", "<<cleared>>", "<<consumption>>", "<<oa_cost>>", "<<discom_cost>>{/monthlyData}"]
        for idx, t in enumerate(tags):
            apply_body_style(new_t1.rows[1].cells[idx], t)
            
        t1._element.addprevious(new_t1._element)
        t1._element.getparent().remove(t1._element)
        
    if t2:
        new_t2 = doc.add_table(rows=2, cols=6)
        new_t2.style = 'Table Grid'
        
        headers = ["Month", "Cleared vs Actual consumption %", "Power Purchase Cost (Discom Only)", "Power Purchase Cost (With Prolt)", "Total Saving", "Saving/Unit"]
        for idx, h in enumerate(headers):
            apply_header_style(new_t2.rows[0].cells[idx], h)
            
        tags = ["{#monthlyData}<<month_name>>", "<<cleared_pct>>", "<<ppc_discom>>", "<<ppc_prolt>>", "<<saving>>", "<<saving_unit>>{/monthlyData}"]
        for idx, t in enumerate(tags):
            apply_body_style(new_t2.rows[1].cells[idx], t)
            
        t2._element.addprevious(new_t2._element)
        t2._element.getparent().remove(t2._element)
        
    doc.save(filepath)

# First reset them from git so we have the original formatting
import subprocess
subprocess.run(['git', 'checkout', 'backend/assets/templates/technical_proposal_template.docx'])
subprocess.run(['git', 'checkout', 'backend/assets/templates/commercial_proposal_template.docx'])

fix_doc('backend/assets/templates/technical_proposal_template.docx')
fix_doc('backend/assets/templates/commercial_proposal_template.docx')
