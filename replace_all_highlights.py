import sys
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def to_camel_case(text):
    # Remove special chars and keep only alphanumeric and spaces
    clean = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    words = clean.split()
    if not words: return "tag"
    return words[0].lower() + ''.join(w.capitalize() for w in words[1:])

def replace_highlights(filepath):
    doc = Document(filepath)
    
    mapping = {
        'Chaurishi Ayurveda LLP, Noida': '<<clientName>>, <<address>>',
        'Chaurishi': '<<clientName>>',
        'Chaurishi Ayurveda LLP': '<<clientName>>',
        'Chaurishi Ayurveda LLP’s': '<<clientName>>’s',
        'Chaurishi Ayurveda LLP’': '<<clientName>>’',
        'Chaurishi Ayurveda': '<<clientName>>',
        '2,11,034': '<<monthlySavings>>',
        '25,32,409': '<<totalSavings>>',
        '150': '<<paybackDays>>',
        'Ayurveda LLP, Noida': '<<industryName>>, <<address>>',
        'January': '<<billMonth>>',
        'January 2026': '<<billMonthYear>>',
        'May 2026': '<<currentMonthYear>>',
        'May  2026': '<<currentMonthYear>>',
    }

    def process_runs(runs):
        for r in runs:
            if r.font.highlight_color is not None and r.font.highlight_color != WD_COLOR_INDEX.AUTO:
                original_text = r.text
                stripped = original_text.strip()
                if not stripped:
                    r.font.highlight_color = None
                    continue
                    
                if stripped in mapping:
                    r.text = r.text.replace(stripped, mapping[stripped])
                else:
                    # Generic fallback: make a tag out of the text if it's numbers or short
                    if len(stripped) < 20:
                        tag_name = to_camel_case(stripped)
                        r.text = r.text.replace(stripped, f'<<{tag_name}>>')
                    else:
                        # For long sentences, just make a generic tag
                        tag_name = "longTextTag"
                        r.text = r.text.replace(stripped, f'<<{tag_name}>>')
                
                r.font.highlight_color = None

    for p in doc.paragraphs:
        process_runs(p.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_runs(p.runs)
                    
    # Handle the specific table with Sanctioned Load
    for table in doc.tables:
        if len(table.rows) >= 2:
            row0_texts = [c.text.strip() for c in table.rows[0].cells]
            if len(row0_texts) >= 3 and 'Sanctioned Load' in row0_texts[0] and 'Connectivity' in row0_texts[1] and 'DISCOM Name' in row0_texts[2]:
                table.rows[1].cells[0].text = '<<sanctionedLoadKw>> kW'
                table.rows[1].cells[1].text = '<<voltageLevel>>'
                table.rows[1].cells[2].text = '<<discom>>'
                if len(row0_texts) > 3 and 'Feeder Type' in row0_texts[3]:
                    table.rows[1].cells[3].text = '<<industryName>>'
                    
                # Remove highlights in headers
                for cell in table.rows[0].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.highlight_color = None

    doc.save(filepath)
    print(f"Processed {filepath}")

replace_highlights("backend/assets/templates/technical_proposal_template.docx")
replace_highlights("backend/assets/templates/commercial_proposal_template.docx")
