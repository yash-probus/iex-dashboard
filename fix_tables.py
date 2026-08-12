from docx import Document
import re
import os

def replace_in_doc(filepath):
    doc = Document(filepath)
    
    # Let's directly iterate over paragraphs and tables to replace EXACT match tags
    # Wait, in the docx, the tags are currently like <<103255>>.
    # We will replace <<103255>> with <<m1_cleared>>, etc.
    
    replacements = {
        '<<103255>>': '<<m1_cleared>>',
        '<<94063>>': '<<m2_cleared>>',
        '<<85298>>': '<<m3_cleared>>',
        '<<88932>>': '<<m4_cleared>>',
        '<<75678>>': '<<m5_cleared>>',
        '<<100690>>': '<<m6_cleared>>',
        
        '<<105107>>': '<<m1_consumption>>',
        '<<101142>>': '<<m2_consumption>>',
        '<<86184>>': '<<m3_consumption>>',
        '<<102300>>': '<<m4_consumption>>',
        '<<89976>>': '<<m5_consumption>>',
        '<<101370>>': '<<m6_consumption>>',
        
        '<<610282>>': '<<m1_oa_cost>>',
        '<<551964>>': '<<m2_oa_cost>>',
        '<<502401>>': '<<m3_oa_cost>>',
        '<<477565>>': '<<m4_oa_cost>>',
        '<<429705>>': '<<m5_oa_cost>>',
        '<<455197>>': '<<m6_oa_cost>>',
        
        '<<815506>>': '<<m1_discom_cost>>',
        '<<735555>>': '<<m2_discom_cost>>',
        '<<670614>>': '<<m3_discom_cost>>',
        '<<711324>>': '<<m4_discom_cost>>',
        '<<583239>>': '<<m5_discom_cost>>',
        '<<777080>>': '<<m6_discom_cost>>',
        
        '<<98>>': '<<m1_cleared_pct>>',
        '<<93>>': '<<m2_cleared_pct>>',
        # '<<99>>' appears twice! For Feb and May!
        # We can't safely replace <<99>> with just m3_cleared_pct. 
        # But wait! Let's check how many times <<99>> is in the tables. 
        # Actually, let's just use string replace on the text runs directly.
        
        '<<790>>': '<<m1_ppc_discom>>',
        '<<782>>': '<<m2_ppc_discom>>',
        '<<786>>': '<<m3_ppc_discom>>',
        '<<800>>': '<<m4_ppc_discom>>',
        '<<771>>': '<<m5_ppc_discom>>',
        '<<772>>': '<<m6_ppc_discom>>',
        
        '<<591>>': '<<m1_ppc_prolt>>',
        '<<587>>': '<<m2_ppc_prolt>>',
        '<<589>>': '<<m3_ppc_prolt>>',
        '<<537>>': '<<m4_ppc_prolt>>',
        '<<568>>': '<<m5_ppc_prolt>>',
        '<<452>>': '<<m6_ppc_prolt>>',
        
        '<<205224>>': '<<m1_saving>>',
        '<<183590>>': '<<m2_saving>>',
        '<<168214>>': '<<m3_saving>>',
        '<<233760>>': '<<m4_saving>>',
        '<<153534>>': '<<m5_saving>>',
        '<<321884>>': '<<m6_saving>>',
        
        '<<199>>': '<<m1_saving_unit>>',
        '<<195>>': '<<m2_saving_unit>>',
        '<<197>>': '<<m3_saving_unit>>',
        '<<263>>': '<<m4_saving_unit>>',
        '<<203>>': '<<m5_saving_unit>>',
        '<<320>>': '<<m6_saving_unit>>',
        
        # Months in headers
        "Dec'25": "<<m1_name>>",
        "Jan'26": "<<m2_name>>",
        "Feb'26": "<<m3_name>>",
        "March'26": "<<m4_name>>",
        "Apr'26": "<<m5_name>>",
        "May'26": "<<m6_name>>",
    }
    
    # We must iterate cell by cell if we want to distinguish the two <<99>> occurrences.
    # To be safe, we will just iterate over all tables and do global replace for unique ones.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)
                                
                        # Handle the <<99>> manually by looking at the column context?
                        # Since <<99>> is used for Feb'26 and May'26 cleared %, it's best to just let it be <<m3_cleared_pct>> for both if we use global replace.
                        # Wait! In the table, row 3 cell 3 is Feb'26 (99), row 3 cell 6 is May'26 (99).
                        if '<<99>>' in run.text:
                            # We can just check the run's text context or just leave it for now.
                            # I will replace it with <<m3_cleared_pct>> for now, and the frontend will map both if needed.
                            # Actually, it's better to just do global replace. The table structure is static.
                            pass

    # A more robust way to handle duplicate <<99>>:
    # Just map <<99>> to <<m3_cleared_pct>> globally. If there are multiple, they all get the same tag.
    # Then in frontend, we can map `m3_cleared_pct` and `m6_cleared_pct` to the same tag if we have to, 
    # but wait, they won't be the same for actual client data!
    # If the template says <<m3_cleared_pct>>, it will render Month 3's data for Month 6! This is a bug.
    
    doc.save(filepath)

# Let's use a smarter approach: we know exactly which table and cell.
def fix_precise(filepath):
    doc = Document(filepath)
    # Table index 0 might not be it. Let's find the tables.
    
    for table in doc.tables:
        # Check if this is Table 1 (Metric | Dec'25 | Jan'26 ...)
        if len(table.rows) >= 4 and len(table.rows[0].cells) >= 6:
            if 'Metric' in table.rows[0].cells[0].text:
                # Update headers
                table.rows[0].cells[1].text = '<<m1_name>>'
                table.rows[0].cells[2].text = '<<m2_name>>'
                table.rows[0].cells[3].text = '<<m3_name>>'
                table.rows[0].cells[4].text = '<<m4_name>>'
                table.rows[0].cells[5].text = '<<m5_name>>'
                if len(table.rows[0].cells) > 6:
                    table.rows[0].cells[6].text = '<<m6_name>>'
                
                # Row 1: Total Cleared Units
                table.rows[1].cells[1].text = '<<m1_cleared>>'
                table.rows[1].cells[2].text = '<<m2_cleared>>'
                table.rows[1].cells[3].text = '<<m3_cleared>>'
                table.rows[1].cells[4].text = '<<m4_cleared>>'
                table.rows[1].cells[5].text = '<<m5_cleared>>'
                if len(table.rows[1].cells) > 6:
                    table.rows[1].cells[6].text = '<<m6_cleared>>'
                    
                # Row 2: Total Consumption
                table.rows[2].cells[1].text = '<<m1_consumption>>'
                table.rows[2].cells[2].text = '<<m2_consumption>>'
                table.rows[2].cells[3].text = '<<m3_consumption>>'
                table.rows[2].cells[4].text = '<<m4_consumption>>'
                table.rows[2].cells[5].text = '<<m5_consumption>>'
                if len(table.rows[2].cells) > 6:
                    table.rows[2].cells[6].text = '<<m6_consumption>>'
                    
                # Row 3: OA Cost
                table.rows[3].cells[1].text = '<<m1_oa_cost>>'
                table.rows[3].cells[2].text = '<<m2_oa_cost>>'
                table.rows[3].cells[3].text = '<<m3_oa_cost>>'
                table.rows[3].cells[4].text = '<<m4_oa_cost>>'
                table.rows[3].cells[5].text = '<<m5_oa_cost>>'
                if len(table.rows[3].cells) > 6:
                    table.rows[3].cells[6].text = '<<m6_oa_cost>>'
                    
                # Row 4: Discom Cost
                table.rows[4].cells[1].text = '<<m1_discom_cost>>'
                table.rows[4].cells[2].text = '<<m2_discom_cost>>'
                table.rows[4].cells[3].text = '<<m3_discom_cost>>'
                table.rows[4].cells[4].text = '<<m4_discom_cost>>'
                table.rows[4].cells[5].text = '<<m5_discom_cost>>'
                if len(table.rows[4].cells) > 6:
                    table.rows[4].cells[6].text = '<<m6_discom_cost>>'

        # Check if this is Table 2 (Month | Cleared vs Actual ...)
        if len(table.rows) >= 6 and len(table.rows[0].cells) >= 6:
            if 'Month' in table.rows[0].cells[0].text and 'Cleared vs' in table.rows[0].cells[1].text:
                # Row 1: m1
                table.rows[1].cells[0].text = '<<m1_name>>'
                table.rows[1].cells[1].text = '<<m1_cleared_pct>>'
                table.rows[1].cells[2].text = '<<m1_ppc_discom>>'
                table.rows[1].cells[3].text = '<<m1_ppc_prolt>>'
                table.rows[1].cells[4].text = '<<m1_saving>>'
                table.rows[1].cells[5].text = '<<m1_saving_unit>>'
                
                # Row 2: m2
                table.rows[2].cells[0].text = '<<m2_name>>'
                table.rows[2].cells[1].text = '<<m2_cleared_pct>>'
                table.rows[2].cells[2].text = '<<m2_ppc_discom>>'
                table.rows[2].cells[3].text = '<<m2_ppc_prolt>>'
                table.rows[2].cells[4].text = '<<m2_saving>>'
                table.rows[2].cells[5].text = '<<m2_saving_unit>>'
                
                # Row 3: m3
                table.rows[3].cells[0].text = '<<m3_name>>'
                table.rows[3].cells[1].text = '<<m3_cleared_pct>>'
                table.rows[3].cells[2].text = '<<m3_ppc_discom>>'
                table.rows[3].cells[3].text = '<<m3_ppc_prolt>>'
                table.rows[3].cells[4].text = '<<m3_saving>>'
                table.rows[3].cells[5].text = '<<m3_saving_unit>>'
                
                # Row 4: m4
                table.rows[4].cells[0].text = '<<m4_name>>'
                table.rows[4].cells[1].text = '<<m4_cleared_pct>>'
                table.rows[4].cells[2].text = '<<m4_ppc_discom>>'
                table.rows[4].cells[3].text = '<<m4_ppc_prolt>>'
                table.rows[4].cells[4].text = '<<m4_saving>>'
                table.rows[4].cells[5].text = '<<m4_saving_unit>>'
                
                # Row 5: m5
                table.rows[5].cells[0].text = '<<m5_name>>'
                table.rows[5].cells[1].text = '<<m5_cleared_pct>>'
                table.rows[5].cells[2].text = '<<m5_ppc_discom>>'
                table.rows[5].cells[3].text = '<<m5_ppc_prolt>>'
                table.rows[5].cells[4].text = '<<m5_saving>>'
                table.rows[5].cells[5].text = '<<m5_saving_unit>>'
                
                # Row 6: m6
                if len(table.rows) > 6:
                    table.rows[6].cells[0].text = '<<m6_name>>'
                    table.rows[6].cells[1].text = '<<m6_cleared_pct>>'
                    table.rows[6].cells[2].text = '<<m6_ppc_discom>>'
                    table.rows[6].cells[3].text = '<<m6_ppc_prolt>>'
                    table.rows[6].cells[4].text = '<<m6_saving>>'
                    table.rows[6].cells[5].text = '<<m6_saving_unit>>'

    doc.save(filepath)

fix_precise('backend/assets/templates/technical_proposal_template.docx')
fix_precise('backend/assets/templates/commercial_proposal_template.docx')
