import docx
import sys

def extract(path):
    doc = docx.Document(path)
    highlights = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                highlights.add(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            highlights.add(run.text)
    for h in highlights:
        print(f'"{h}"')

extract(sys.argv[1])
