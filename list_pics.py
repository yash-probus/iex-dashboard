from docx import Document

def list_pics(filepath):
    doc = Document(filepath)
    print(f"--- {filepath} ---")
    pic_index = 0
    for i, p in enumerate(doc.paragraphs):
        has_pic = False
        for r in p.runs:
            if '<w:drawing>' in r._element.xml:
                has_pic = True
        if has_pic:
            pic_index += 1
            print(f"Picture {pic_index}:")
            # Print text of previous 3 paragraphs and next 3 paragraphs
            start = max(0, i-3)
            end = min(len(doc.paragraphs), i+4)
            for j in range(start, end):
                prefix = "  -> " if j == i else "     "
                print(f"{prefix}P{j}: {doc.paragraphs[j].text.strip()}")
            print()

list_pics('backend/assets/templates/commercial_proposal_template.docx')
list_pics('backend/assets/templates/technical_proposal_template.docx')
