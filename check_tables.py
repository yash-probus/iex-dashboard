import re
from docx import Document

doc = Document('backend/assets/templates/technical_proposal_template.docx')
text = ""
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text += p.text + "\n"

# Find all << and >>
starts = [m.start() for m in re.finditer('<<', text)]
ends = [m.start() for m in re.finditer('>>', text)]

print(f"Starts: {len(starts)}, Ends: {len(ends)}")

# Let's extract all tags
tags = re.findall(r'<<(.*?)>>', text)
print("Tags in tables:", tags)
