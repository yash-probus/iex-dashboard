import sys
import re
from docx import Document

doc = Document(sys.argv[1])
text = ""
for p in doc.paragraphs:
    text += p.text + "\n"
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text += p.text + "\n"

tags = re.findall(r'<<(.*?)>>', text)
print("Tags:", tags)
