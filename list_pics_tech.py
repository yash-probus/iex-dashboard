from docx import Document

def list_pics(filepath):
    doc = Document(filepath)
    print(f"--- {filepath} ---")
    pic_index = 0
    for i, p in enumerate(doc.paragraphs):
        has_pic = False
        for r in p.runs:
            if '<w:drawing>' in r._element.xml:
                has_pic = True
        if has_pic:
            pic_index += 1
            if 'DETAILED' in doc.paragraphs[max(0, i-6)].text or 'DETAILED' in doc.paragraphs[max(0, i-10)].text or 'Rajeev' in doc.paragraphs[max(0, i-5)].text:
                print(f"Picture {pic_index}:")
                start = max(0, i-6)
                end = min(len(doc.paragraphs), i+4)
                for j in range(start, end):
                    prefix = "  -> " if j == i else "     "
                    print(f"{prefix}P{j}: {doc.paragraphs[j].text.strip()}")
                print()

list_pics('backend/assets/templates/technical_proposal_template.docx')
