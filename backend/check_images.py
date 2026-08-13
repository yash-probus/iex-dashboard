from docx import Document
import sys

doc = Document('backend/assets/templates/technical_proposal_template.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")

for i, p in enumerate(doc.paragraphs):
    if '<<' in p.text:
        print(f"P {i}: {p.text}")
    for run in p.runs:
        # Check for drawing elements (images)
        if 'graphic' in run._element.xml:
            print(f"P {i}: Found an image/graphic here!")

