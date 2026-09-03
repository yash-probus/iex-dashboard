import sys
import json
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_rupee(amount):
    try:
        if amount is None or str(amount).strip() == '':
            return '₹ 0'
        # Parse the amount safely, removing commas, Rs, etc.
        amount_str = str(amount).replace(',', '').replace('₹', '').replace('Rs.', '').replace('Rs', '').strip()
        num = float(amount_str)
        # Format as Indian Rupee
        s, *d = str(num).partition(".")
        r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
        return f"₹ {r}"
    except (ValueError, TypeError):
        return str(amount)

def safe_float(val, default=0.0):
    try:
        if val is None or str(val).strip() == '':
            return default
        val_str = str(val).replace(',', '').replace('₹', '').replace('Rs.', '').replace('Rs', '').strip()
        return float(val_str)
    except (ValueError, TypeError):
        return default

def generate_proposal(data_json):
    if data_json.endswith('.json') and os.path.isfile(data_json):
        with open(data_json, 'r') as f:
            data = json.load(f)
    else:
        data = json.loads(data_json)
    
    template_path = data.get('template_path', '/Users/yashgupta/IEX-Dashboard/commercial_proposal_template.docx')
    output_path = data.get('output_path', '/Users/yashgupta/IEX-Dashboard/generated_proposal.docx')
    
    doc = docx.Document(template_path)
    
    # 1. Update Title Paragraph P2
    client_name = data.get('client_name') or data.get('industry_name') or 'CLIENT'
    connectivity_str = str(data.get('connectivity') or data.get('voltage_level') or '11 kV')
    
    def remove_highlight(run):
        rPr = run._element.rPr
        if rPr is not None:
            highlights = rPr.xpath('./w:highlight')
            for hl in highlights:
                rPr.remove(hl)

    def replace_in_text(text):
        if not text: return text
        if 'XXXXXXXXXXXXX' in text:
            text = text.replace('XXXXXXXXXXXXX', client_name.upper())
        if '33KV' in text:
            text = text.replace('33KV', connectivity_str.upper())
        if '33 KV' in text:
            text = text.replace('33 KV', connectivity_str.upper())
        if '33kV' in text:
            text = text.replace('33kV', connectivity_str)
        return text

    for p in doc.paragraphs:
        for run in p.runs:
            remove_highlight(run)
            if run.text:
                run.text = replace_in_text(run.text)
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        remove_highlight(run)
                        if run.text:
                            run.text = replace_in_text(run.text)

    # Helper function to set cell text with centered alignment and bold formatting
    def set_cell_value(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    # 2. Update Table 0 (Facility Parameters: Sanctioned Load, Connectivity, DISCOM Name, Feeder Type)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        if len(t0.rows) > 1:
            r1 = t0.rows[1].cells
            sanctioned_load = str(data.get('sanctioned_load') or data.get('sanctioned_load_kw') or '1000 kW')
            if not sanctioned_load.lower().endswith('kw') and not sanctioned_load.lower().endswith('kva'):
                sanctioned_load += ' kW'
            set_cell_value(r1[0], sanctioned_load, bold=True)
            set_cell_value(r1[1], str(data.get('connectivity') or data.get('voltage_level') or '11 kV'), bold=True)
            set_cell_value(r1[2], str(data.get('discom_name') or data.get('discom') or 'DISCOM'), bold=True)
            set_cell_value(r1[3], str(data.get('feeder_type') or 'Dedicated Feeder'), bold=True)

    # 3. Update Table 3 (Pricing/BOQ Table)
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        
        # Calculate totals
        c_supply = safe_float(data.get('abt_supply_cost'), 450000)
        c_service = safe_float(data.get('abt_service_cost'), 350000)
        c_liaison = safe_float(data.get('utility_liaisoning_cost'), 300000)
        c_total_abt = c_supply + c_service + c_liaison
        
        c_bg = safe_float(data.get('bank_guarantee_cost'), 150000)
        
        c_iex = safe_float(data.get('iex_annual_fee'), 100000)
        c_noc = safe_float(data.get('sldc_monthly_noc'), 7000)
        c_st11 = safe_float(data.get('st11_settlement'), 20000)
        c_total_recurring = c_iex + c_noc + c_st11
        
        tm = str(data.get('trading_margin', '2p/kWh'))
        pf = str(data.get('platform_fee', '2p/kWh'))
        vs = str(data.get('value_share', '15%'))
        if not vs.endswith('%'): vs += '%'
        smart = safe_float(data.get('smart_metering_infra'), 125000)

        # Row 1: ABT Metering Total
        if len(t3.rows) > 1: set_cell_value(t3.rows[1].cells[-1], format_rupee(c_total_abt), bold=True)
        
        # Row 2: Supply
        if len(t3.rows) > 2: 
            supply_cell = t3.rows[2].cells[1]
            indoor_count = data.get('indoor_ctpt_count')
            if indoor_count:
                for p in supply_cell.paragraphs:
                    p.text = p.text.replace('1 Nos.', indoor_count)
            if data.get('remove_outdoor_supply'):
                for p in supply_cell.paragraphs:
                    if 'Outdoor CT/PT' in p.text:
                        p.text = ''
            set_cell_value(t3.rows[2].cells[-1], format_rupee(c_supply))
            
        # Row 3: Service
        if len(t3.rows) > 3: set_cell_value(t3.rows[3].cells[-1], format_rupee(c_service))
        
        # Row 4: Liaisoning
        if len(t3.rows) > 4: set_cell_value(t3.rows[4].cells[-1], format_rupee(c_liaison))
        
        # Row 5: SLDC Security Deposit Total
        if len(t3.rows) > 5: set_cell_value(t3.rows[5].cells[-1], format_rupee(c_bg), bold=True)
        
        # Row 6: UPSLDC Bank Guarantee
        if len(t3.rows) > 6: set_cell_value(t3.rows[6].cells[-1], format_rupee(c_bg))
        
        # Row 7: Fixed Recurring Charges Total
        if len(t3.rows) > 7: set_cell_value(t3.rows[7].cells[-1], format_rupee(c_total_recurring), bold=True)
        
        # Row 8: IEX Annual
        if len(t3.rows) > 8: set_cell_value(t3.rows[8].cells[-1], format_rupee(c_iex))
        
        # Row 9: NOC
        if len(t3.rows) > 9: set_cell_value(t3.rows[9].cells[-1], format_rupee(c_noc))
        
        # Row 10: ST-11
        if len(t3.rows) > 10: set_cell_value(t3.rows[10].cells[-1], format_rupee(c_st11))
        
        # Row 12: Trading Margin
        if len(t3.rows) > 12: set_cell_value(t3.rows[12].cells[-1], tm)
        
        # Row 13: Platform Fee
        if len(t3.rows) > 13: set_cell_value(t3.rows[13].cells[-1], pf)
        
        # Row 14: Value Share
        if len(t3.rows) > 14: set_cell_value(t3.rows[14].cells[-1], vs)
        
        # Row 15: Smart Metering
        if len(t3.rows) > 15: set_cell_value(t3.rows[15].cells[-1], format_rupee(smart))

    # 4. Update Table 4 (Terms & Conditions)
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        smart_payment_term = str(data.get('smart_metering_infra_payment_term', '100% Advance against PO/PI'))
        
        # Searching for 'Prolt Energy Smart Metering Infra' in the second column to replace the third column
        for row in t4.rows:
            if len(row.cells) > 1 and 'Prolt Energy Smart Metering' in row.cells[1].text:
                set_cell_value(row.cells[2], smart_payment_term, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
                break

    # 5. Format the last page (remove header, set 0 margins, stretch image)
    try:
        import copy
        from docx.oxml.ns import qn
        from docx.shared import Cm
        
        body = doc._element.body
        doc_sectPr = body.sectPr
        
        sectPr1 = copy.deepcopy(doc_sectPr)
        
        img_idx = -1
        for i in range(len(doc.paragraphs)-1, -1, -1):
            if list(doc.paragraphs[i]._element.iter(qn('w:drawing'))):
                img_idx = i
                break
                
        if img_idx > 0:
            p_prev = doc.paragraphs[img_idx - 1]._element
            pPr = p_prev.get_or_add_pPr()
            if pPr.find(qn('w:sectPr')) is None:
                pPr.append(sectPr1)
            
            sec2 = doc.sections[-1]
            sec2.top_margin = Cm(0)
            sec2.bottom_margin = Cm(0)
            sec2.left_margin = Cm(0)
            sec2.right_margin = Cm(0)
            
            for header_ref in doc_sectPr.findall(qn('w:headerReference')):
                doc_sectPr.remove(header_ref)
            for footer_ref in doc_sectPr.findall(qn('w:footerReference')):
                doc_sectPr.remove(footer_ref)
                
            img_para = doc.paragraphs[img_idx]._element
            for drawing in img_para.iter(qn('w:drawing')):
                wp_anchor = drawing.find(qn('wp:anchor'))
                if wp_anchor is not None:
                    ext = wp_anchor.find(qn('wp:extent'))
                    if ext is not None:
                        ext.set('cx', '7562448') 
                        ext.set('cy', '10689336')
                        
                    for xfrm in wp_anchor.xpath('.//*[local-name()="xfrm"]'):
                        ext2 = xfrm.xpath('./*[local-name()="ext"]')
                        if ext2:
                            ext2[0].set('cx', '7562448')
                            ext2[0].set('cy', '10689336')
    except Exception as e:
        print(f"Warning: Could not format last page - {e}")

    if os.path.dirname(output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"Successfully generated commercial proposal at {output_path}")

if __name__ == '__main__':
    if len(sys.argv) > 1:
        generate_proposal(sys.argv[1])
    else:
        print("No input json provided")
