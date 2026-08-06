import docx

def extract_highlighted_text(docx_path):
    doc = docx.Document(docx_path)
    highlighted_texts = []
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                highlighted_texts.append({
                    'text': run.text,
                    'color': run.font.highlight_color
                })
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            highlighted_texts.append({
                                'text': run.text,
                                'color': run.font.highlight_color
                            })
                            
    return highlighted_texts

if __name__ == "__main__":
    path = "/Users/yashgupta/IEX-Dashboard/DRAFT TECHNICAL PROPOSAL_V2.docx"
    highlights = extract_highlighted_text(path)
    
    print(f"Found {len(highlights)} highlighted items:")
    for h in highlights:
        # print only if color is YELLOW (7) or we can print all
        if h['text'].strip():
            print(f"[{h['color']}] {h['text'].strip()}")
