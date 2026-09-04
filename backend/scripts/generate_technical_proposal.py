import sys
import json
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
import io
import base64
import copy
import calendar

def get_png_dimensions(data):
    if len(data) >= 24 and data.startswith(b'\x89PNG\r\n\x1a\n'):
        import struct
        width, height = struct.unpack('>II', data[16:24])
        return width, height
    return None

def safe_float(val):
    if val is None:
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    try:
        cleaned = str(val).replace('₹', '').replace(',', '').replace('Rs.', '').replace('Rs', '').replace('%', '').replace('/', '').strip()
        if not cleaned:
            return 0.0
        return float(cleaned)
    except Exception:
        return 0.0

def format_rupee(amount, decimals=0):
    try:
        if amount is None or str(amount).strip() == '':
            return '₹ 0'
        amount_str = str(amount).replace(',', '').replace('₹', '').replace('Rs.', '').replace('Rs', '').strip()
        num = float(amount_str)
        
        if decimals == 0:
            s = str(int(round(num)))
            d = ""
        else:
            s, _, d = f"{num:.{decimals}f}".partition(".")
            
        # Indian numbering system formatting
        if len(s) > 3:
            last_three = s[-3:]
            remaining = s[:-3]
            groups = []
            while remaining:
                groups.append(remaining[-2:])
                remaining = remaining[:-2]
            groups.reverse()
            formatted_s = ",".join(groups) + "," + last_three
        else:
            formatted_s = s
            
        if d:
            return f"₹ {formatted_s}.{d}"
        else:
            return f"₹ {formatted_s}"
    except (ValueError, TypeError):
        return str(amount)

def replace_in_text(text, replacements):
    if not text: return text
    for old_val, new_val in replacements.items():
        if old_val in text:
            text = text.replace(old_val, str(new_val))
    return text

def duplicate_row(row):
    tr = row._tr
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)
    return docx.table._Row(new_tr, row._parent)

def generate_proposal(data_json):
    if data_json.endswith('.json') and os.path.isfile(data_json):
        with open(data_json, 'r') as f:
            data = json.load(f)
    else:
        data = json.loads(data_json)
    
    template_path = data.get('template_path', '/Users/yashgupta/IEX-Dashboard/TECHNICAL PROPOSAL_11KV.docx')
    output_path = data.get('output_path', '/Users/yashgupta/IEX-Dashboard/generated_technical_proposal.docx')
    
    doc = docx.Document(template_path)
    client_name = data.get('client_name') or data.get('industry_name') or 'CLIENT XYZ'
    
    # Sanitize load to avoid double KVA
    load_val = str(data.get('sanctioned_load') or '1000').strip()
    if not load_val.upper().endswith('KVA'):
        sanctioned_load = f"{load_val} KVA"
    else:
        sanctioned_load = load_val.upper()

    # Sanitize connectivity to avoid double KV
    conn_val = str(data.get('connectivity') or '11').strip()
    if not conn_val.upper().endswith('KV'):
        connectivity = f"{conn_val} KV"
    else:
        connectivity = conn_val.upper()
    discom_name = data.get('discom_name') or 'PUVVNL'
    feeder_type = data.get('feeder_type') or 'Dedicated'
    
    avg_monthly_savings = safe_float(data.get('average_monthly_savings', 211034))
    avg_annual_savings = safe_float(data.get('average_annual_savings', 2532409))

    # 1. Process Monthly Data for Dynamic Replacement Metrics
    monthly_data = data.get('monthlyData', [])
    total_months = len(monthly_data)
    positive_months = 0
    oa_shares = []
    savings_list = []
    savings_per_units = []
    
    strongest_saving = 0.0
    strongest_month_label = "May 2026"
    
    for item in monthly_data:
        savings = safe_float(item.get('savings', 0))
        savings_list.append(savings)
        if savings > 0:
            positive_months += 1
            
        total_energy = safe_float(item.get('total_energy_kwh') or item.get('discom_only', {}).get('volume', 1))
        market_energy = safe_float(item.get('consumer_bus_units') or item.get('total_market_energy_kwh') or 0)
        oa_share = (market_energy / total_energy * 100) if total_energy > 0 else 0
        if oa_share > 100:
            oa_share = 100
        oa_shares.append(oa_share)
        
        savings_pu = safe_float(item.get('savings_per_unit', 0))
        savings_per_units.append(savings_pu)
        
        if savings > strongest_saving:
            strongest_saving = savings
            month_str = item.get('month', '')
            if '-' in month_str:
                parts = month_str.split('-')
                if len(parts) == 2:
                    try:
                        year = int(parts[0])
                        month_num = int(parts[1])
                        month_names_full = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
                        strongest_month_label = f"{month_names_full[month_num - 1]} {year}"
                    except Exception:
                        pass
            elif '/' in month_str:
                parts = month_str.split('/')
                if len(parts) == 2:
                    try:
                        month_num = int(parts[0])
                        year = int(parts[1])
                        month_names_full = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
                        strongest_month_label = f"{month_names_full[month_num - 1]} {year}"
                    except Exception:
                        pass

    # Aggregates
    min_oa = min(oa_shares) if oa_shares else 0
    max_oa = max(oa_shares) if oa_shares else 0
    min_saving = min(savings_list) if savings_list else 0
    max_saving = max(savings_list) if savings_list else 0
    min_pu = min(savings_per_units) if savings_per_units else 0
    max_pu = max(savings_per_units) if savings_per_units else 0
    
    total_market_all = sum(safe_float(item.get('consumer_bus_units') or item.get('total_market_energy_kwh') or 0) for item in monthly_data)
    total_energy_all = sum(safe_float(item.get('total_energy_kwh') or item.get('discom_only', {}).get('volume', 1)) for item in monthly_data)
    avg_oa_share = (total_market_all / total_energy_all * 100) if total_energy_all > 0 else 0
    if avg_oa_share > 100:
        avg_oa_share = 100
    avg_oa_str = f"Approximately {round(avg_oa_share)}%"

    # Format numbers into visual proposal metrics
    positive_months_ratio = f"{positive_months} / {total_months}" if total_months > 0 else "0 / 0"
    oa_share_range = f"{round(min_oa)}% - {round(max_oa)}%"
    savings_range_lakhs = f"₹{(min_saving/100000):.2f}L - ₹{(max_saving/100000):.2f}L"
    saving_per_unit_range = f"₹{min_pu:.2f} - ₹{max_pu:.2f}"
    
    num_to_words = {
        1: "One Month", 2: "Two Months", 3: "Three Months", 4: "Four Months", 
        5: "Five Months", 6: "Six Months", 7: "Seven Months", 8: "Eight Months", 
        9: "Nine Months", 10: "Ten Months", 11: "Eleven Months", 12: "Twelve Months"
    }
    total_months_words = num_to_words.get(total_months, f"{total_months} Months")

    # Calculate Metering Charge Payback
    total_cost = safe_float(data.get('smart_metering_infra_cost', 0))
    monthly_probus_savings = round(safe_float(data.get('monthly_probus_savings_amount', 0)))
    payback_str = "150 days"
    if monthly_probus_savings > 0 and total_cost > 0:
        months_val = total_cost / monthly_probus_savings
        full_months = int(months_val)
        days = round((months_val - full_months) * 30)
        payback_str = f"{full_months} months and {days} days"

    # Expand replacements dictionary to handle curly apostrophes and specific character lengths
    replacements = {
        'Approximately 95%': avg_oa_str,
        '95%': f"{round(avg_oa_share)}%",
        '150 days': payback_str,
        'XXXXXXXXXXXXXX': client_name.upper(),
        'XXXXXXXXXXXXX': client_name.upper(),
        'XXXXXXXXXXXX': client_name.upper(),
        'XXXXXXXXXXX': client_name.upper(),
        'XXXXXXXXXX': client_name.upper(),
        'Client XYZ': client_name.upper(),
        'CLIENT XYZ': client_name.upper(),
        'Chaurishi Ayurveda LLP’s': f"{client_name}’s",
        'Chaurishi Ayurveda LLP’': f"{client_name}’",
        'Chaurishi Ayurveda LLP\'s': f"{client_name}'s",
        'Chaurishi Ayurveda LLP': client_name,
        '1000 KVA': sanctioned_load,
        '11 KV': connectivity,
        'PUVVNL': discom_name,
        'Dedicated': feeder_type,
        '₹2,11,034': format_rupee(avg_monthly_savings, decimals=0),
        '₹25,32,409': format_rupee(avg_annual_savings, decimals=0),
        '₹3,21,884': format_rupee(strongest_saving, decimals=0),
        'May 2026': strongest_month_label,
        '6 / 6': positive_months_ratio,
        '84% - 99%': oa_share_range,
        '₹1.55L - ₹3.22L': savings_range_lakhs,
        '₹1.95 - ₹3.20': saving_per_unit_range,
        '₹25 lakh': f"₹{round(avg_annual_savings / 100000)} lakh",
        '25 lakh': f"{round(avg_annual_savings / 100000)} lakh",
        'Six Months': total_months_words,
        'Six months': total_months_words,
        'six-month': f"{total_months}-month",
        'six month': f"{total_months} month"
    }
    
    # 2. Update text paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = replace_in_text(run.text, replacements)
                
    # 3. Update tables (except the savings table which we'll handle separately)
    savings_table_index = -1
    for i, t in enumerate(doc.tables):
        is_savings_table = False
        if len(t.rows) > 0 and len(t.rows[0].cells) > 2:
            if 'Bill Months' in t.rows[0].cells[0].text or 'Billing Period' in t.rows[0].cells[1].text:
                savings_table_index = i
                is_savings_table = True
        
        if not is_savings_table:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for run in p.runs:
                            if run.text:
                                run.text = replace_in_text(run.text, replacements)
    
    # 4. Handle savings table
    if savings_table_index != -1 and monthly_data:
        t = doc.tables[savings_table_index]
        template_row_idx = 3
        
        if len(t.rows) > template_row_idx:
            # 4.1 Delete original data rows below the template row first
            for i in range(len(t.rows) - 1, template_row_idx, -1):
                row = t.rows[i]
                t._element.remove(row._tr)
            
            # 4.2 Get a reference to the single remaining template data row
            template_row = t.rows[template_row_idx]
            
            # 4.3 Generate new rows based on monthly_data
            for item in monthly_data:
                new_row = duplicate_row(template_row)
                cells = new_row.cells
                
                # Derive month_label and billing_period from month
                month_str = item.get('month', '')
                month_label = month_str
                billing_period = ''
                
                if '-' in month_str:
                    parts = month_str.split('-')
                    if len(parts) == 2:
                        try:
                            year = int(parts[0])
                            month_num = int(parts[1])
                            month_names_short = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
                            month_label = f"{month_names_short[month_num]}'{str(year)[2:]}"
                            last_day = calendar.monthrange(year, month_num)[1]
                            billing_period = f"01/{month_num:02d} - {last_day:02d}/{month_num:02d}"
                        except Exception:
                            pass
                elif '/' in month_str:
                    parts = month_str.split('/')
                    if len(parts) == 2:
                        try:
                            month_num = int(parts[0])
                            year = int(parts[1])
                            month_names_short = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
                            month_label = f"{month_names_short[month_num]}'{str(year)[2:]}"
                            last_day = calendar.monthrange(year, month_num)[1]
                            billing_period = f"01/{month_num:02d} - {last_day:02d}/{month_num:02d}"
                        except Exception:
                            pass

                actual_units = safe_float(item.get('discom_only', {}).get('volume', 0))
                
                # Calculate Cleared OA % using consumer bus units
                consumer_bus_units = safe_float(item.get('consumer_bus_units', 0))
                cleared_oa_pct = (consumer_bus_units / actual_units * 100) if actual_units > 0 else 0
                if cleared_oa_pct > 100:
                    cleared_oa_pct = 100
                cleared_oa = f"{round(cleared_oa_pct)}%"
                
                discom_total = format_rupee(item.get('discom_only', {}).get('total_amount', 0), decimals=0)
                discom_pu = safe_float(item.get('discom_only', {}).get('per_unit_effective', 0))
                mix_total = format_rupee(item.get('oa_mix', {}).get('total_amount', 0), decimals=0)
                mix_pu = safe_float(item.get('oa_mix', {}).get('per_unit_effective', 0))
                savings = format_rupee(item.get('savings', 0), decimals=0)
                savings_pu = safe_float(item.get('savings_per_unit', 0))
                
                if len(cells) > 0:
                    cells[0].text = str(month_label)
                if len(cells) > 1:
                    cells[1].text = str(billing_period)
                if len(cells) > 2:
                    cells[2].text = f"{int(actual_units):,}" if actual_units else "0"
                if len(cells) > 3:
                    cells[3].text = str(cleared_oa)
                if len(cells) > 4:
                    cells[4].text = str(discom_total)
                if len(cells) > 5:
                    cells[5].text = f"₹ {discom_pu:.2f}" if discom_pu else "₹ 0.00"
                if len(cells) > 6:
                    cells[6].text = str(mix_total)
                if len(cells) > 7:
                    cells[7].text = f"₹ {mix_pu:.2f}" if mix_pu else "₹ 0.00"
                if len(cells) > 8:
                    cells[8].text = str(savings)
                if len(cells) > 9:
                    cells[9].text = f"₹ {savings_pu:.2f}" if savings_pu else "₹ 0.00"
                    
                # Center-align text and style in savings cells
                for cell_idx in range(len(cells)):
                    p = cells[cell_idx].paragraphs[0] if cells[cell_idx].paragraphs else cells[cell_idx].add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Aptos'
                        run.font.size = Pt(9)
                            
            # 4.4 Delete the original template row itself
            t._element.remove(template_row._tr)
                
    # 5. Handle Chart Replacement (Targeted to the correct savings section chart)
    chart_b64 = data.get('monthly_savings_chart')
    if chart_b64:
        try:
            image_data = base64.b64decode(chart_b64)
            image_stream = io.BytesIO(image_data)
            
            # Find the savings section header first
            savings_header_idx = -1
            for idx, p in enumerate(doc.paragraphs):
                if 'Consistent Across the Billing History' in p.text or 'SAVINGS FOR YOUR BUSINESS' in p.text:
                    savings_header_idx = idx
                    break
            
            chart_replaced = False
            if savings_header_idx != -1:
                # Search forward from the header for the first paragraph containing a graphic
                for idx in range(savings_header_idx, len(doc.paragraphs)):
                    p = doc.paragraphs[idx]
                    if 'Graphic' in p._element.xml or 'pic:pic' in p._element.xml:
                        p.clear()
                        run = p.add_run()
                        run.add_picture(image_stream, width=Inches(6.0))
                        chart_replaced = True
                        print(f"Replaced monthly savings chart at paragraph {idx}")
                        break
            
            # Fallback if not found via search
            if not chart_replaced:
                for idx, p in enumerate(doc.paragraphs):
                    if idx in [110, 109, 111, 112]:
                        if 'Graphic' in p._element.xml or 'pic:pic' in p._element.xml:
                            p.clear()
                            run = p.add_run()
                            run.add_picture(image_stream, width=Inches(6.0))
                            print(f"Replaced monthly savings chart at paragraph {idx} (fallback)")
                            break
        except Exception as e:
            print(f"Error replacing image: {e}")

    # 5.1 Handle Dashboard Hero/KPIs Screenshot replacement (image14.png under '8. Savings Summary Snapshot')
    hero_kpi_b64 = data.get('dashboard_hero_kpis_screenshot') or data.get('dashboard_screenshot')
    if hero_kpi_b64:
        try:
            image_data = base64.b64decode(hero_kpi_b64)
            # Find the paragraph containing "8. Savings Summary Snapshot"
            snapshot_idx = -1
            for idx, p in enumerate(doc.paragraphs):
                if '8. Savings Summary Snapshot' in p.text or 'Savings Summary Snapshot' in p.text:
                    snapshot_idx = idx
                    break
            
            if snapshot_idx != -1:
                # Search forward from snapshot header for the first paragraph containing a drawing
                for idx in range(snapshot_idx, len(doc.paragraphs)):
                    p = doc.paragraphs[idx]
                    if 'w:drawing' in p._element.xml:
                        import re
                        rIds = re.findall(r'r:embed=\"([^\"]+)\"', p._element.xml)
                        if rIds:
                            rId = rIds[0]
                            if rId in doc.part.related_parts:
                                part = doc.part.related_parts[rId]
                                part._blob = image_data
                                print(f"Successfully replaced dashboard hero/kpis image part {part.partname} at paragraph {idx}")
                                
                                # Resize image to maintain aspect ratio
                                dims = get_png_dimensions(image_data)
                                if dims:
                                    new_w_px, new_h_px = dims
                                    aspect_ratio = new_h_px / new_w_px
                                    target_width_inches = 6.5
                                    target_height_inches = target_width_inches * aspect_ratio
                                    width_emu = int(target_width_inches * 914400)
                                    height_emu = int(target_height_inches * 914400)
                                    
                                    for run in p.runs:
                                        drawing = run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                        if drawing is not None:
                                            for elem in drawing.iter():
                                                if elem.get('cx') is not None:
                                                    elem.set('cx', str(width_emu))
                                                if elem.get('cy') is not None:
                                                    elem.set('cy', str(height_emu))
                                            print(f"Resized hero/kpis image to {target_width_inches}x{target_height_inches:.2f} inches")
                                break
        except Exception as e:
            print(f"Error replacing hero/kpis screenshot: {e}")

    # Clear all remaining yellow highlights in the document (from both paragraphs and tables)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.highlight_color = None
            
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = None

    if os.path.dirname(output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated technical proposal at {output_path}")

if __name__ == '__main__':
    if len(sys.argv) > 1:
        generate_proposal(sys.argv[1])
    else:
        print("No input json provided")
