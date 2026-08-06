import os
import docx

def replace_in_run(run, mapping):
    # Sort keys by length descending to replace longer matches first
    for old in sorted(mapping.keys(), key=len, reverse=True):
        new = mapping[old]
        if old in run.text:
            run.text = run.text.replace(old, new)
            # Remove highlight after replacing to clean up the template
            run.font.highlight_color = None
            
def generate_commercial_template(input_path, output_path):
    doc = docx.Document(input_path)
    
    # Basic mapping based on extracted highlights
    mapping = {
        "XXXXXXXXXXXXX": "<<client_name>>",
        "Connectivity": "<<connectivity>>",
        "Feeder Type": "<<feeder_type>>",
        "Sanctioned Load": "<<sanctioned_load>>",
        "DISCOM Name": "<<discom_name>>",
        "₹450,000.00": "<<total_cost>>",
        "₹350,000.00": "<<subtotal>>",
        "₹300,000.00": "<<software_cost>>",
        "1,50,000.00": "<<hardware_cost>>",
        "20,000.00": "<<recurring_cost>>",
        "100% Advance against PO/PI": "<<payment_terms>>",
        "2p/kWh": "<<success_fee>>",
        "15%": "<<savings_percentage>>"
    }
    
    # Simple run-level replacement
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                replace_in_run(run, mapping)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            replace_in_run(run, mapping)
                            
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Commercial Template saved to {output_path}")

if __name__ == "__main__":
    input_doc = "/Users/yashgupta/IEX-Dashboard/COMMERCIAL PROPOSAL_V2 (1).docx"
    output_doc = "/Users/yashgupta/IEX-Dashboard/backend/assets/templates/commercial_template.docx"
    generate_commercial_template(input_doc, output_doc)
