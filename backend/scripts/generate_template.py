import os
import docx

def replace_in_run(run, mapping):
    # Sort keys by length descending to replace longer matches first
    for old in sorted(mapping.keys(), key=len, reverse=True):
        new = mapping[old]
        if old in run.text:
            run.text = run.text.replace(old, new)
            # Remove highlight after replacing to clean up the template
            run.font.highlight_color = None
            
def generate_template(input_path, output_path):
    doc = docx.Document(input_path)
    
    # Basic mapping based on extracted highlights
    mapping = {
        "Chaurishi Ayurveda LLP": "<<client_name>>",
        "Chaurishi Ayurveda": "<<client_short_name>>",
        "Noida": "<<client_city>>",
        "Chaurishi": "<<client_short_name>>",
        "May 2026": "<<end_month>>",
        "January 2026": "<<start_month>>",
        "Six months": "<<duration>>",
        "of January and May 2026": "of <<start_month_name>> and <<end_month_name>> <<year>>",
        "2,11,034": "<<total_units>>",
        "25,32,409": "<<total_savings>>",
        "150": "<<sanctioned_load>>",
        "98%": "<<savings_percentage>>",
        "₹7.90": "<<utility_rate>>",
        "₹5.91": "<<iex_rate>>",
        "₹2,05,224": "<<monthly_savings>>",
        "₹1.99": "<<rate_diff>>"
    }
    
    # Simple run-level replacement
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                replace_in_run(run, mapping)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            replace_in_run(run, mapping)
                            
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template saved to {output_path}")

if __name__ == "__main__":
    input_doc = "/Users/yashgupta/IEX-Dashboard/DRAFT TECHNICAL PROPOSAL_V2.docx"
    output_doc = "/Users/yashgupta/IEX-Dashboard/backend/assets/templates/proposal_template.docx"
    generate_template(input_doc, output_doc)
