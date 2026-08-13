from docx import Document

doc = Document('backend/assets/templates/technical_proposal_template.docx')

# Helper function to remove drawing elements
def clear_graphics(p):
    for run in p.runs:
        for element in run._element:
            if 'graphic' in element.tag or 'drawing' in element.tag:
                run._element.remove(element)
    
    # Check fallback directly in paragraph
    for element in p._element:
        if 'graphic' in element.tag or 'drawing' in element.tag:
            p._element.remove(element)

# Clear graphics from the Insight sections
for i in range(85, 100):
    clear_graphics(doc.paragraphs[i])

for i in range(104, 110):
    clear_graphics(doc.paragraphs[i])

# Insert the tags where the main images were
# P 86 was the first image
doc.paragraphs[86].text = "<<%first_insight_screenshot>>"

# P 106 was the second image
doc.paragraphs[106].text = "<<%second_insight_screenshot>>"

doc.save('backend/assets/templates/technical_proposal_template.docx')
print("Done!")
