from docx import Document

doc = Document('backend/assets/templates/technical_proposal_template.docx')
doc2 = Document('backend/assets/templates/commercial_proposal_template.docx')

def process_paragraph(p):
    if "of January and May 2026" in p.text:
        for run in p.runs:
            if run.text == "of January and May 2026":
                run.text = "of <<first_insight_month>> and <<second_insight_month>>"
    
    if "January 2026 Insights" in p.text:
        for run in p.runs:
            if "January 2026 Insights" in run.text:
                run.text = run.text.replace("January 2026 Insights", "<<first_insight_month>> Insights")

    if "May  2026" in p.text or "May 2026" in p.text:
        for run in p.runs:
            if "May  2026" in run.text:
                run.text = run.text.replace("May  2026", "<<second_insight_month>>")
            elif "May 2026" in run.text:
                run.text = run.text.replace("May 2026", "<<second_insight_month>>")

    if "17%" in p.text:
        for run in p.runs:
            if "17%" in run.text:
                run.text = run.text.replace("17%", "<<first_insight_reduction>>%")

    if "2,00,000" in p.text:
        for run in p.runs:
            if "2,00,000" in run.text:
                run.text = run.text.replace("2,00,000", "<<first_insight_saving>>")

    if "January 2026" in p.text and "Insights" not in p.text:
        for run in p.runs:
            if "January 2026" in run.text:
                run.text = run.text.replace("January 2026", "<<first_insight_month>>")
                
    if "January" in p.text:
        for run in p.runs:
            if "January" in run.text:
                run.text = run.text.replace("January", "<<first_insight_month_only>>")

    if "May" in p.text and "May 2026" not in p.text and "May  2026" not in p.text:
        for run in p.runs:
            if "May" in run.text:
                run.text = run.text.replace("May", "<<second_insight_month_only>>")

for p in doc.paragraphs:
    process_paragraph(p)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)

doc.save('backend/assets/templates/technical_proposal_template.docx')

for p in doc2.paragraphs:
    process_paragraph(p)
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)
doc2.save('backend/assets/templates/commercial_proposal_template.docx')
print("Done!")
