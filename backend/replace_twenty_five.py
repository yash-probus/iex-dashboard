import sys
from docx import Document

doc = Document('backend/assets/templates/technical_proposal_template.docx')

def process_paragraph(p):
    if "Twenty Five lakhs" in p.text:
        # We know Run 1 has 'Twenty Five' and Run 2 has ' lakhs'
        for i, run in enumerate(p.runs):
            if run.text == 'Twenty Five':
                run.text = '<<savings_in_words>>'
                print("Replaced Run 1")
            elif run.text == ' lakhs':
                run.text = ''
                print("Cleared Run 2")

for p in doc.paragraphs:
    process_paragraph(p)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)

doc.save('backend/assets/templates/technical_proposal_template.docx')
print("Saved changes!")
