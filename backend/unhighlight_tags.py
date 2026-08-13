import sys
from docx import Document

def unhighlight_tags(doc_path):
    print(f"Processing {doc_path}...")
    doc = Document(doc_path)
    
    def process_runs(runs):
        in_tag = False
        for run in runs:
            # Check if this run starts a tag
            if "<<" in run.text:
                in_tag = True
            
            # If we are inside a tag (or this run is a tag itself), remove highlight
            if in_tag or "<<" in run.text or ">>" in run.text:
                if run.font.highlight_color is not None:
                    run.font.highlight_color = None
                    print(f"Unhighlighted: {run.text.strip()}")
            
            # Check if this run ends a tag
            if ">>" in run.text:
                in_tag = False

    for p in doc.paragraphs:
        process_runs(p.runs)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_runs(p.runs)
                    
    doc.save(doc_path)
    print(f"Saved {doc_path}")

unhighlight_tags('backend/assets/templates/technical_proposal_template.docx')
unhighlight_tags('backend/assets/templates/commercial_proposal_template.docx')
