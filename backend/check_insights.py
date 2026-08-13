from docx import Document

doc = Document('backend/assets/templates/technical_proposal_template.docx')

def process_paragraph(p):
    if "Insights" in p.text or "January" in p.text:
        print(f"P: {p.text}")
        for i, run in enumerate(p.runs):
            print(f"  Run {i}: '{run.text}', Highlight: {run.font.highlight_color}")

for p in doc.paragraphs:
    process_paragraph(p)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)
