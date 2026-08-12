import sys
from docx import Document

def modify(filepath, output_path):
    doc = Document(filepath)
    for p in doc.paragraphs:
        for r in p.runs:
            text = r.text
            if 'Chaurishi Ayurveda LLP, Noida' in text:
                r.text = text.replace('Chaurishi Ayurveda LLP, Noida', '<<clientName>>, <<address>>')
                r.font.highlight_color = None
            elif 'Chaurishi Ayurveda LLP’' in text:
                r.text = text.replace('Chaurishi Ayurveda LLP’', '<<clientName>>’')
                r.font.highlight_color = None
            elif 'Chaurishi Ayurveda LLP' in text:
                r.text = text.replace('Chaurishi Ayurveda LLP', '<<clientName>>')
                r.font.highlight_color = None
            elif 'Chaurishi Ayurveda' in text:
                r.text = text.replace('Chaurishi Ayurveda', '<<clientName>>')
                r.font.highlight_color = None
            elif 'Chaurishi' in text:
                r.text = text.replace('Chaurishi', '<<clientName>>')
                r.font.highlight_color = None

    # For tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Let's check for the header row with Sanctioned Load
                if 'Sanctioned Load' in cell.text:
                    # Let's just find the cell index
                    pass

                for p in cell.paragraphs:
                    for r in p.runs:
                        text = r.text
                        if 'Chaurishi Ayurveda LLP, Noida' in text:
                            r.text = text.replace('Chaurishi Ayurveda LLP, Noida', '<<clientName>>, <<address>>')
                            r.font.highlight_color = None
                        elif 'Chaurishi Ayurveda LLP’' in text:
                            r.text = text.replace('Chaurishi Ayurveda LLP’', '<<clientName>>’')
                            r.font.highlight_color = None
                        elif 'Chaurishi Ayurveda LLP' in text:
                            r.text = text.replace('Chaurishi Ayurveda LLP', '<<clientName>>')
                            r.font.highlight_color = None
                        elif 'Chaurishi Ayurveda' in text:
                            r.text = text.replace('Chaurishi Ayurveda', '<<clientName>>')
                            r.font.highlight_color = None
                        elif 'Chaurishi' in text:
                            r.text = text.replace('Chaurishi', '<<clientName>>')
                            r.font.highlight_color = None
                        elif '2,11,034' in text:
                            r.text = text.replace('2,11,034', '<<monthlySavings>>')
                            r.font.highlight_color = None
                        elif '25,32,409' in text:
                            r.text = text.replace('25,32,409', '<<totalSavings>>')
                            r.font.highlight_color = None
                        elif '150' in text and len(text.strip()) == 3:
                            r.text = text.replace('150', '<<paybackDays>>')
                            r.font.highlight_color = None

    # Handle the specific table with Sanctioned Load
    for table in doc.tables:
        if len(table.rows) >= 2:
            row0_texts = [c.text.strip() for c in table.rows[0].cells]
            if 'Sanctioned Load' in row0_texts[0] and 'Connectivity' in row0_texts[1] and 'DISCOM Name' in row0_texts[2]:
                table.rows[1].cells[0].text = '<<sanctionedLoadKw>> kW'
                table.rows[1].cells[1].text = '<<voltageLevel>>'
                table.rows[1].cells[2].text = '<<discom>>'
                if len(row0_texts) > 3 and 'Feeder Type' in row0_texts[3]:
                    table.rows[1].cells[3].text = '<<industryName>>'
                    
                # Remove highlights in headers
                for cell in table.rows[0].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.highlight_color = None

    doc.save(output_path)
    print(f"Saved {output_path}")

if __name__ == "__main__":
    modify(sys.argv[1], sys.argv[2])
