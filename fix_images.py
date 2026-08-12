from docx import Document
from docx.shared import Pt
import re

def fix_doc(filepath):
    doc = Document(filepath)
    
    # Replace "Twenty Five lakhs" with "<<savings_in_words>>"
    for p in doc.paragraphs:
        if 'Twenty Five lakhs' in p.text:
            p.text = p.text.replace('Twenty Five lakhs', '<<savings_in_words>>')
            
            # Make the word '<<savings_in_words>>' have a yellow highlight if possible, or bold
            for r in p.runs:
                if '<<savings_in_words>>' in r.text:
                    r.font.bold = True
                    # Setting highlight color might be complex, let's just make it bold

    # Replace the graph picture
    # In docxtemplater, image tags are {%tagname}
    # We find the paragraph before the "You are losing around" text or just insert the tag in the first paragraph of the page
    # Let's find the paragraph that contains the inline shape (picture)
    
    for p in doc.paragraphs:
        # Find runs with pictures
        has_picture = False
        for r in p.runs:
            if '<w:drawing>' in r._element.xml:
                has_picture = True
                
        # If the paragraph before it contains "MONTHLY SAVINGS OPPORTUNITY", we know this is the chart image!
        # But wait, the chart image itself contains the title "MONTHLY SAVINGS OPPORTUNITY" embedded IN the image.
        # Let's check the paragraph BEFORE it.
        # A simpler way: Just delete all pictures that are charts and insert {%monthly_savings_chart}
    
    # We know the chart is after "DETAILED SAVINGS ANALYSIS" and before "You are losing around".
    found_detailed = False
    for p in doc.paragraphs:
        if 'DETAILED SAVINGS ANALYSIS' in p.text:
            found_detailed = True
        
        if found_detailed:
            for r in p.runs:
                if '<w:drawing>' in r._element.xml:
                    # We found the picture!
                    # Let's clear the paragraph and put {%monthly_savings_chart}
                    p.clear()
                    p.add_run('{%monthly_savings_chart}')
                    found_detailed = False # Stop replacing pictures
                    break
        
    doc.save(filepath)

# reset
import subprocess
subprocess.run(['git', 'checkout', '3b7207cc', '--', 'backend/assets/templates/technical_proposal_template.docx'])
subprocess.run(['git', 'checkout', '3b7207cc', '--', 'backend/assets/templates/commercial_proposal_template.docx'])

fix_doc('backend/assets/templates/technical_proposal_template.docx')
fix_doc('backend/assets/templates/commercial_proposal_template.docx')
