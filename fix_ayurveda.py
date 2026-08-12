from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def fix(filepath):
    doc = Document(filepath)
    for p in doc.paragraphs:
        for r in p.runs:
            if 'Ayurveda LLP' in r.text or 'Ayurveda' in r.text:
                if r.font.highlight_color is not None:
                    r.text = r.text.replace(' Ayurveda LLP, Noida', ' <<industryName>>')
                    r.text = r.text.replace('Ayurveda LLP, Noida', '<<industryName>>')
                    r.font.highlight_color = None
    
    # Also check tables just in case
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if 'Ayurveda LLP' in r.text or 'Ayurveda' in r.text:
                            if r.font.highlight_color is not None:
                                r.text = r.text.replace(' Ayurveda LLP, Noida', ' <<industryName>>')
                                r.text = r.text.replace('Ayurveda LLP, Noida', '<<industryName>>')
                                r.font.highlight_color = None
                                
    doc.save(filepath)

fix('backend/assets/templates/technical_proposal_template.docx')
fix('backend/assets/templates/commercial_proposal_template.docx')
