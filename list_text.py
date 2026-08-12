from docx import Document

def list_text(filepath):
    doc = Document(filepath)
    print(f"--- {filepath} ---")
    for i, p in enumerate(doc.paragraphs):
        if 'dashboard' in p.text.lower() or '{%' in p.text:
            print(f"P{i}: {p.text}")
    for i, t in enumerate(doc.tables):
        for row in t.rows:
            for cell in row.cells:
                if 'dashboard' in cell.text.lower() or '{%' in cell.text:
                    print(f"Table {i}: {cell.text}")

list_text('backend/assets/templates/commercial_proposal_template.docx')
list_text('backend/assets/templates/technical_proposal_template.docx')
