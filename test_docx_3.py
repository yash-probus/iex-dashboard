import docx
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

doc = docx.Document('backend/assets/templates/COMMERCIAL PROPOSAL_33KV.docx')

try:
    body = doc._element.body
    doc_sectPr = body.sectPr
    
    sectPr1 = copy.deepcopy(doc_sectPr)
    
    img_idx = -1
    for i in range(len(doc.paragraphs)-1, -1, -1):
        if list(doc.paragraphs[i]._element.iter(qn('w:drawing'))):
            img_idx = i
            break
            
    if img_idx > 0:
        p_prev = doc.paragraphs[img_idx - 1]._element
        pPr = p_prev.get_or_add_pPr()
        if pPr.find(qn('w:sectPr')) is None:
            pPr.append(sectPr1)
        
        sec2 = doc.sections[-1]
        sec2.top_margin = Cm(0)
        sec2.bottom_margin = Cm(0)
        sec2.left_margin = Cm(0)
        sec2.right_margin = Cm(0)
        
        for header_ref in doc_sectPr.findall(qn('w:headerReference')):
            doc_sectPr.remove(header_ref)
        for footer_ref in doc_sectPr.findall(qn('w:footerReference')):
            doc_sectPr.remove(footer_ref)
            
        img_para = doc.paragraphs[img_idx]._element
        for drawing in img_para.iter(qn('w:drawing')):
            wp_anchor = drawing.find(qn('wp:anchor'))
            if wp_anchor is not None:
                ext = wp_anchor.find(qn('wp:extent'))
                if ext is not None:
                    ext.set('cx', '7562448') 
                    ext.set('cy', '10689336')
                    
                for xfrm in wp_anchor.xpath('.//*[local-name()="xfrm"]'):
                    ext2 = xfrm.xpath('./*[local-name()="ext"]')
                    if ext2:
                        ext2[0].set('cx', '7562448')
                        ext2[0].set('cy', '10689336')

    doc.save('backend/assets/templates/test_output_2.docx')
    print("Success")
except Exception as e:
    print(f"Error: {e}")

